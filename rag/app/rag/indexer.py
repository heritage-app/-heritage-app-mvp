"""
Document indexing module using LlamaIndex.
Handles one-time indexing and document processing.
"""

import logging
from pathlib import Path
from typing import List, Optional
from llama_index.core import VectorStoreIndex, Document
from llama_index.core.node_parser import SentenceSplitter, MarkdownNodeParser
import fitz  # PyMuPDF for PDF files
import re
import io
import pandas as pd
import docx
import json

from app.rag.embeddings import get_embeddings
from app.rag.vector_store import get_vector_store, collection_exists, get_qdrant_client
from app.storage.service import StorageService
from app.storage.supabase_client import get_supabase
from app.rag.constants import COLLECTION_NAME, COLLECTION_MAP, get_collection_name

logger = logging.getLogger(__name__)


class BibleRefiner:
    """
    Unified extraction and refinement engine for Bible documents.
    Normalizes PDF, DOCX, and unstructured text into high-fidelity Verse Records.
    """
    
    BOOK_MAPPING = {
        "genesis": "Mose klɛŋklɛŋ wolo",
        "exodus": "Mose wolo ni ji enyɔ",
        "leviticus": "Mose wolo ni ji etɛ",
        "numbers": "Mose wolo ni ji ejwɛ",
        "deuteronomy": "Mose wolo ni ji enumɔ",
        "joshua": "Yoshua wolo",
        "judges": "Kojolɔi awolo",
        "ruth": "Rut wolo",
        "psalms": "Lalafo Wolo",
        "proverbs": "Abɛi awolo"
    }

    @classmethod
    def _num_to_ga(cls, n: int) -> str:
        from app.rag.utils import num_to_ga
        return num_to_ga(n)

    @classmethod
    def _num_to_en(cls, n: int) -> str:
        """Simple English number-to-word for chapter titles."""
        units = {1: "One", 2: "Two", 3: "Three", 4: "Four", 5: "Five", 
                 6: "Six", 7: "Seven", 8: "Eight", 9: "Nine"}
        teens = {11: "Eleven", 12: "Twelve", 13: "Thirteen", 14: "Fourteen", 15: "Fifteen",
                 16: "Sixteen", 17: "Seventeen", 18: "Eighteen", 19: "Nineteen"}
        tens = {10: "Ten", 20: "Twenty", 30: "Thirty", 40: "Forty", 50: "Fifty", 
                60: "Sixty", 70: "Seventy", 80: "Eighty", 90: "Ninety"}
        if n in units: return units[n]
        if n in teens: return teens[n]
        if n in tens: return tens[n]
        if n < 100: return f"{tens[n // 10 * 10]} {units[n % 10]}"
        return str(n)

    @classmethod
    def get_ga_label(cls, num: int) -> str:
        """Converts a digit into a 'Kuku ni ji [word]' label."""
        return f"Kuku ni ji {cls._num_to_ga(num)}"

    @classmethod
    def get_ga_chapter_title(cls, num: int) -> str:
        """Converts a digit into a 'Yitso [Word]' title."""
        word = cls._num_to_ga(num)
        # Capitalize first word
        if " " in word:
            parts = word.split(" ")
            word = f"{parts[0].capitalize()} {' '.join(parts[1:])}"
        else:
            word = word.capitalize()
        return f"Yitso {word}"

    @classmethod
    def get_trad_book(cls, book_name: str) -> str:
        return cls.BOOK_MAPPING.get(book_name.lower(), book_name)

    @classmethod
    def _build_bible_node_text(cls, r: dict) -> str:
        """Unified high-fidelity text template for indexing Bible verses."""
        ga_name = r.get("ga_version_name") or "Ŋmalɛ Krɔŋkrɔŋ Lɛ"
        ga_abbr = r.get("ga_version_abbr") or "NEGAB"
        en_name = r.get("english_version_name") or "King James Version"
        en_abbr = r.get("english_version_abbr") or "KJV"
        
        # Citation header (reference_display)
        # Format: Mose klɛŋklɛŋ wolo, Yitso 50, Kuku ni ji 1 (Genesis 50:1)
        ref_display = r.get("reference_display")
        if not ref_display:
            trad_book = r.get("traditional_book") or cls.get_trad_book(r.get("book", ""))
            ga_ch = r.get("chapter_title_ga") or cls.get_ga_chapter_title(r.get("chapter_num", 0))
            ga_v = r.get("ga_label") or cls.get_ga_label(r.get("verse_num", 0))
            v_ref = r.get("verse_ref") or f"{r.get('book')} {r.get('chapter_num')}:{r.get('verse_num')}"
            ref_display = f"{trad_book}, {ga_ch}, {ga_v} ({v_ref})"
            r["reference_display"] = ref_display # Save back to record
        
        # Ensure other requested fields are in record
        r["ga_version_name"] = ga_name
        r["ga_version_abbr"] = ga_abbr
        r["english_version_name"] = en_name
        r["english_version_abbr"] = en_abbr

        # Build node content for LLM extraction
        node_text = (
            f"METADATA:\n"
            f"reference_display: {ref_display}\n"
            f"verse_ref: {r.get('verse_ref')}\n"
            f"ga_version_name: {ga_name}\n"
            f"ga_version_abbr: {ga_abbr}\n"
            f"english_version_name: {en_name}\n"
            f"english_version_abbr: {en_abbr}\n"
            f"source_name: {r.get('source_name', 'Bible Archive')}\n\n"
            f"Ga Content:\n{r.get('ga', '').strip()}\n\n"
            f"English Content:\n{r.get('en', '').strip()}"
        )
        return node_text

    @classmethod
    def parse_unstructured(cls, text: str, filename: str) -> List[dict]:
        """
        Robustly parses multi-chapter interleaved Ga/English blocks.
        Supports patterns like (Ga 17, En 17, Ga 18, En 18).
        """
        # 0. Pre-cleaning: Remove noise markers without deleting valid lines
        # Aggressive removal of footers and copyright info
        noise_phrases = [
            "Currently Selected:", 
            "Bible in Ga Language", 
            "© Bible Society of Ghana",
            "Learn More", 
            "Rights in the",
            "©"
        ]
        clean_text = text
        for noise in noise_phrases:
            # Case insensitive removal, replacing with a single space to avoid concatenating words
            clean_text = re.sub(re.escape(noise), " ", clean_text, flags=re.IGNORECASE)
        
        # 1. Identify all Chapter Headers (Broad, anchorless search for multi-book support)
        # Supports phonetic variants of Genesis (Jɛnɛsis, JƐNƐSIS, JEnesis, etc.)
        # Now handles "Genesis 17", "Genesis Chapter 17", "Jɛnɛsis Yitso 17", etc.
        header_pattern = re.compile(
            r'(Genesis|J[ɛ\u025b\u0190eE]n[ɛ\u025b\u0190eE]sis|Exodus|Psalms|Mose|Job|Matthew|Mark|Luke|John|Revelation)'
            r'\s+(?:Chapter|Yitso|YITS\u0196|Yitso\s+ni\s+ji)?' # Optional chapter keywords
            r'\s*(\d+)', 
            re.IGNORECASE
        )
        matches = list(header_pattern.finditer(clean_text))
        
        if not matches:
            # Absolute fallback for direct split (try to find the first book name)
            parts = re.split(r'(?i)\bGenesis\b.*?\d+', clean_text, maxsplit=1)
            if len(parts) < 2: return []
            ga_block, en_block = parts[0], parts[1]
            ch_map = {0: {"ga": ga_block, "en": en_block, "book": "Genesis"}}
        else:
            # 2. Slice text into blocks
            blocks = []
            for i in range(len(matches)):
                start = matches[i].start()
                end = matches[i+1].start() if i + 1 < len(matches) else len(clean_text)
                book_str = matches[i].group(1)
                ch_num = int(matches[i].group(2))
                is_ga = book_str.lower().startswith('j') or book_str.lower() == "mose"
                content = clean_text[start:end]
                
                blocks.append({
                    "chapter_num": ch_num,
                    "is_ga": is_ga,
                    "text": content,
                    "book": "Exodus" if "exodus" in book_str.lower() else ("Psalms" if "psalms" in book_str.lower() else "Genesis")
                })

            # 3. Group by Chapter Number
            ch_map = {}
            for b in blocks:
                c_num = b["chapter_num"]
                if c_num not in ch_map:
                    ch_map[c_num] = {"ga": "", "en": "", "book": b["book"]}
                if b["is_ga"]: ch_map[c_num]["ga"] += b["text"] + "\n"
                else: ch_map[c_num]["en"] += b["text"] + "\n"

        # 4. Process Each Chapter
        def extract_verses(block):
            # Improved regex: Matches digit(s) followed by optional space, 
            # and takes content until the next digit(s) that are followed by a non-digit character.
            # This handles "1  Be ni" (Ga) and "1And when" (En) styles.
            pattern = re.compile(r'(\d+)\s*(.*?)(?=\s*\d+[^\d]|$)', re.DOTALL)
            found = pattern.findall(block)
            
            data = {}
            # Match "Kuku ni ji [any word]" at the start (case insensitive, including Unicode)
            label_prefix_pattern = re.compile(r'^\s*Kuku\s+ni\s+ji\s+[\w\u025b\u0254\u014b]+\s*', re.IGNORECASE)
            
            for m in found:
                v_num = int(m[0])
                v_text = m[1].strip()
                # Strip the traditional label "Kuku ni ji..." if it exists to avoid redundancy
                clean_text = label_prefix_pattern.sub("", v_text).strip()
                data[v_num] = clean_text
            return data

        all_records = []
        for ch_num, data in sorted(ch_map.items()):
            ga_data = extract_verses(data["ga"])
            en_data = extract_verses(data["en"])
            
            book = data["book"]
            trad_book = cls.get_trad_book(book)
            ch_title_ga = cls.get_ga_chapter_title(ch_num) if ch_num > 0 else "Bible Archive"
            
            ch_verses = sorted(list(set(ga_data.keys()) | set(en_data.keys())))
            for v in ch_verses:
                all_records.append({
                    "id": f"bible:{book.lower()}:{ch_num}:{v}",
                    "category": "bible",
                    "source_name": filename.replace('_', ' ').split('.')[0],
                    "book": book,
                    "traditional_book": trad_book,
                    "chapter_num": ch_num,
                    "chapter_ref": f"Chapter {ch_num}",
                    "chapter_title_ga": ch_title_ga,
                    "chapter_title_en": f"Chapter {cls._num_to_en(ch_num)}",
                    "section_ga": "",
                    "section_en": "",
                    "verse_num": v,
                    "verse_ref": f"{book} {ch_num}:{v}",
                    "ga_verse_label": cls.get_ga_label(v),
                    "ga": ga_data.get(v, ""),
                    "en": en_data.get(v, ""),
                    "ga_version_abbr": "NEGAB",  # Default if not provided
                    "english_version_abbr": "KJV" # Default if not provided
                })
        return all_records

    @classmethod
    def extract_raw_text(cls, content: bytes, filename: str) -> str:
        """Extracts UTF-8 text from bytes based on file extension."""
        ext = filename.lower().split('.')[-1]
        try:
            if ext == 'pdf':
                doc = fitz.open(stream=content, filetype="pdf")
                text = ""
                for page in doc:
                    text += page.get_text()
                return text
            elif ext == 'docx':
                doc = docx.Document(io.BytesIO(content))
                return "\n".join([p.text for p in doc.paragraphs])
            else:
                return content.decode('utf-8', errors='ignore')
        except Exception:
            return content.decode('utf-8', errors='ignore')

    @classmethod
    def get_refinement_preview(cls, content: bytes, filename: str) -> dict:
        """Returns raw text and refined records for frontend preview."""
        raw_text = cls.extract_raw_text(content, filename)
        records = cls.parse_unstructured(raw_text, filename)
        
        # Build JSONL preview
        jsonl_lines = [json.dumps(r, ensure_ascii=False) for r in records[:5]]
        jsonl_preview = "\n".join(jsonl_lines)
        if len(records) > 5:
            jsonl_preview += f"\n... and {len(records) - 5} more verses"

        return {
            "raw_text": raw_text[:10000],  # Limit preview size
            "refined_records": records,
            "jsonl_preview": jsonl_preview,
            "stats": {
                "verse_count": len(records),
                "chapters": list(set(r["chapter_num"] for r in records))
            }
        }


async def index_document_from_storage(file_path: str, metadata: Optional[dict] = None) -> None:
    """Index a single document directly from Supabase Storage."""
    content = await StorageService.download_file(file_path)
    await index_from_bytes(content, file_path, metadata)


async def index_refined_records(records: List[dict], metadata: Optional[dict] = None) -> None:
    """
    Directly index pre-refined Bible records into the vector store.
    Bypasses extraction and parsing.
    """
    documents = []
    for r in records:
        node_text = BibleRefiner._build_bible_node_text(r)
        
        doc = Document(
            text=node_text,
            metadata={
                **(metadata or {}),
                **r
            }
        )
        documents.append(doc)

    if not documents:
        return

    # Use the unified indexer logic
    VectorStoreIndex.from_documents(
        documents,
        storage_context=get_vector_store(COLLECTION_NAME),
        embed_model=get_embeddings(),
        show_progress=False
    )


async def index_from_bytes(content: bytes, file_path: str, metadata: Optional[dict] = None) -> None:
    """
    Index a document directly from bytes in memory.
    Now with automated Bible Refinement pipeline.
    """
    file_path_obj = Path(file_path)
    file_extension = file_path_obj.suffix.lower()
    filename = file_path_obj.name
    
    text = ""
    # 1. Extraction Stage
    if file_extension == ".pdf":
        doc = fitz.open(stream=content, filetype="pdf")
        for page in doc: text += page.get_text()
        doc.close()
    elif file_extension in [".txt", ".md", ".json", ".jsonl"]:
        text = content.decode('utf-8', errors='ignore')
    elif file_extension == ".docx":
        doc_obj = docx.Document(io.BytesIO(content))
        text = "\n".join([p.text for p in doc_obj.paragraphs])
    else:
        text = content.decode('utf-8', errors='ignore')

    if not text.strip(): return

    # 2. Classification
    category = _detect_category(text, file_path)
    collection_name = get_collection_name(category)
    base_metadata = metadata.copy() if metadata else {}
    base_metadata.update({"category": category, "file_path": file_path, "filename": filename})
    
    # Metadata fields to exclude from embeddings and LLM context to prevent length errors
    excluded_keys = ["file_path", "filename", "id", "status", "uploaded_at", "source_name"]

    # 3. Specialized Parsing & AUTOMATION HOOK
    documents = []
    if category == "bible":
        if file_extension == ".jsonl":
            documents = _parse_bible_jsonl(text, file_path)
        else:
            logger.info(f"Refining Bible document: {filename}...")
            if file_extension == ".pdf":
                documents = _parse_bible_pdf(text, file_path)
            else:
                records = BibleRefiner.parse_unstructured(text, filename)
                if records:
                    for r in records:
                        node_text = BibleRefiner._build_bible_node_text(r)
                        
                        doc = Document(
                            text=node_text, 
                            metadata={**base_metadata, **r},
                            excluded_embed_metadata_keys=excluded_keys,
                            excluded_llm_metadata_keys=excluded_keys
                        )
                        documents.append(doc)
                    
                    # Save Refined sidecar back to Supabase
                    jsonl_content = "\n".join([json.dumps(r) for r in records])
                    refined_path = f"processed/{filename.split('.')[0]}_refined.jsonl"
                    try:
                        await StorageService.upload_file(jsonl_content.encode('utf-8'), refined_path, "application/jsonl")
                        logger.info(f"Saved refined sidecar: {refined_path}")
                    except Exception as e: logger.warning(f"Failed to save sidecar: {e}")

        if not documents: 
            documents = [Document(
                text=text, 
                metadata=base_metadata,
                excluded_embed_metadata_keys=excluded_keys,
                excluded_llm_metadata_keys=excluded_keys
            )]
    else:
        # --- Non-Bible documents ---
        if file_extension == ".jsonl":
            # Route through the JSONL record-level parser so each record gets its own embedding
            documents = _parse_heritage_jsonl(text, file_path, base_metadata, excluded_keys)
        
        if not documents:
            # Fallback: treat as a single document (PDF, DOCX, TXT, etc.)
            documents = [Document(
                text=text,
                metadata=base_metadata,
                excluded_embed_metadata_keys=excluded_keys,
                excluded_llm_metadata_keys=excluded_keys
            )]

    # 4. Insertion
    for doc in documents:
        doc.metadata.update(base_metadata)
        doc.excluded_embed_metadata_keys = list(set(doc.excluded_embed_metadata_keys + excluded_keys))
        doc.excluded_llm_metadata_keys = list(set(doc.excluded_llm_metadata_keys + excluded_keys))

    embeddings = get_embeddings()
    vector_store = get_vector_store(collection_name)
    
    # Bible: larger chunks to handle verse metadata; JSONL heritage: already 1 doc per record, use small overlap
    if category == "bible":
        chunk_size = 2048
    elif file_extension == ".jsonl":
        # Each doc is already a single record — no need to split further
        chunk_size = 512
    else:
        chunk_size = 1024
    node_parser = SentenceSplitter(chunk_size=chunk_size, chunk_overlap=50)
    index = VectorStoreIndex.from_vector_store(vector_store=vector_store, embed_model=embeddings)
    
    for doc in documents:
        nodes = node_parser.get_nodes_from_documents([doc])
        await index.ainsert_nodes(nodes)


async def is_document_indexed(file_path: str) -> bool:
    """
    Check if a document is already indexed in Qdrant by checking for nodes with matching file_path metadata.
    
    Args:
        file_path: Path to document in storage
        
    Returns:
        bool: True if document is indexed, False otherwise
    """
    try:
        from qdrant_client.models import Filter, FieldCondition, MatchValue
        
        # Get Qdrant client
        client = get_qdrant_client()
        
        # Check if collection exists
        if not collection_exists(COLLECTION_NAME):
            return False
        
        # Search for any points with matching file_path in metadata
        try:
            # Use scroll to check for documents with matching metadata
            # Check for file_path or filename in metadata
            results = client.scroll(
                collection_name=COLLECTION_NAME,
                scroll_filter=Filter(
                    must=[
                        FieldCondition(
                            key="file_path",
                            match=MatchValue(value=file_path)
                        )
                    ]
                ),
                limit=1
            )
            
            # If we found any points, the document is indexed
            if results[0]:  # results is a tuple (points, next_page_offset)
                return len(results[0]) > 0
            
            # Also check for filename in metadata (backward compatibility)
            results = client.scroll(
                collection_name=COLLECTION_NAME,
                scroll_filter=Filter(
                    must=[
                        FieldCondition(
                            key="filename",
                            match=MatchValue(value=file_path)
                        )
                    ]
                ),
                limit=1
            )
            
            return len(results[0]) > 0 if results[0] else False
            
        except Exception:
            # If scroll fails, try a different approach - just check if collection has any points
            # This is a fallback - not as precise but better than nothing
            collection_info = client.get_collection(COLLECTION_NAME)
            # If collection exists and has points, we assume it might be indexed
            # This is not perfect but works as a fallback
            return collection_info.points_count > 0
            
    except Exception:
        # If anything fails, assume not indexed to be safe
        return False


async def index_local_document(file_path: str, metadata: Optional[dict] = None) -> None:
    """
    Index a local document file (for backward compatibility and initial indexing).
    
    Args:
        file_path: Path to local document file
        metadata: Optional metadata to attach to document
    """
    # Read file content
    local_path = Path(file_path)
    content = local_path.read_bytes()
    
    # Index from bytes
    await index_from_bytes(content, file_path, metadata)


async def index_directory(directory_path: str, metadata: Optional[dict] = None) -> None:
    """
    Index all documents in a directory.
    
    Args:
        directory_path: Path to directory containing documents
        metadata: Optional metadata to attach to all documents
    """
    dir_path = Path(directory_path)
    
    # Supported file extensions
    supported_extensions = {".pdf", ".txt", ".docx", ".md", ".html", ".json", ".jsonl", ".csv", ".xlsx", ".xls"}
    
    # Find all supported files
    files = [f for f in dir_path.rglob("*") if f.suffix.lower() in supported_extensions]
    
    for file_path in files:
        await index_local_document(str(file_path), metadata)



def _parse_heritage_jsonl(text: str, file_path: str, base_metadata: dict, excluded_keys: list) -> List[Document]:
    """
    Record-level parser for general heritage JSONL files (phrases, stories, data).
    Creates ONE Document per JSON line so every record gets its own vector embedding.
    This permanently fixes the coarse-chunking problem for all future JSONL uploads.
    """
    documents = []
    for line in text.strip().splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            record = json.loads(line)
        except (json.JSONDecodeError, ValueError):
            continue

        # Build a human-readable node text for embedding
        # Support common schemas: {english, ga}, {en, ga}, {text}, {phrase, translation}, etc.
        english = record.get("english") or record.get("en") or record.get("phrase") or ""
        ga_text = record.get("ga") or record.get("translation") or record.get("ga_text") or ""
        category_tag = record.get("category") or record.get("topic") or ""

        if english and ga_text:
            # Standard phrase record: rich searchable text
            node_text = f"English: {english}\nGa: {ga_text}"
            if category_tag:
                node_text = f"[{category_tag}]\n{node_text}"
        elif english or ga_text:
            node_text = english or ga_text
        else:
            # Unknown schema — stringify the whole record
            node_text = " | ".join(f"{k}: {v}" for k, v in record.items() if isinstance(v, str))

        if not node_text.strip():
            continue

        meta = {**base_metadata, **{k: v for k, v in record.items() if isinstance(v, (str, int, float, bool))}}
        doc = Document(
            text=node_text,
            metadata=meta,
            excluded_embed_metadata_keys=excluded_keys,
            excluded_llm_metadata_keys=excluded_keys
        )
        documents.append(doc)

    logger.info(f"JSONL parser: created {len(documents)} records from {Path(file_path).name}")
    return documents


def _detect_category(text: str, file_path: str) -> str:
    """Auto-detect category based on content and filename."""
    path_lower = file_path.lower()
    text_sample = text[:2000].lower()
    if any(k in path_lower for k in ["bible", "bibele", "genesis", "exodus", "psalms", "j\u025bn\u025bsis"]):
        return "bible"
    if "kuku ni ji" in text_sample or ("chapter" in text_sample and "verse" in text_sample):
        return "bible"
    if any(k in path_lower for k in ["story", "stories", "adese", "folk"]):
        return "stories"
    return "heritage"


def _parse_markdown_bible(text: str, file_path: str) -> List[Document]:
    """Specialized parser for the bilingual Bible Markdown format."""
    lines = text.split('\n')
    documents = []
    current_chapter, current_book = "Unknown Chapter", "Genesis"
    if "genesis" in file_path.lower(): current_book = "Genesis"
    
    chapter_pattern = re.compile(r'^##\s+(Yitso\s+.*?\s*\(Chapter\s+.*?\))', re.IGNORECASE)
    row_pattern = re.compile(r'^\|\s*(\d+)\s*\|(.*?)\|(.*?)\|(.*?)\|')

    for line in lines:
        line = line.strip()
        chapter_match = chapter_pattern.match(line)
        if chapter_match:
            current_chapter = chapter_match.group(1)
            continue
        row_match = row_pattern.match(line)
        if row_match:
            v_num = row_match.group(1).strip()
            record = {
                "book": current_book,
                "chapter_num": current_chapter,
                "verse_num": v_num,
                "ga": row_match.group(4).strip(),
                "en": row_match.group(2).strip(),
                "verse_ref": f"{current_book} {current_chapter}:{v_num}",
                "source_name": Path(file_path).name
            }
            node_text = BibleRefiner._build_bible_node_text(record)
            doc = Document(text=node_text, metadata={**record, "category": "bible", "file_path": file_path, "filename": Path(file_path).name})
            documents.append(doc)
    return documents or [Document(text=text)]


def _parse_bible_pdf(text: str, file_path: str) -> List[Document]:
    """Legacy PDF parser using 'Kuku ni ji' markers."""
    documents = []
    current_book = "Genesis" if "genesis" in file_path.lower() else "Bible"
    num_map = {"ekome": "1", "enyɔ": "2", "etɛ": "3", "ejwɛ": "4", "enumɔ": "5", "ekpaa": "6", "kpawo": "7", "kpaanyɔ": "8", "neehu": "9", "nyɔŋma": "10"}
    
    segments = re.split(r'Kuku\s+ni\s+ji', text, flags=re.IGNORECASE)
    header = segments[0].strip()
    chapter_match = re.search(r'Yitso\s+(.*?)\s*\(Chapter\s+(.*?)\)', header, re.IGNORECASE)
    
    def resolve(raw):
        m = re.search(r'\d+', raw)
        if m: return m.group(0)
        for w, v in num_map.items():
            if w in raw.lower(): return v
        return "0"

    curr_ch = resolve(chapter_match.group(2) if chapter_match else "")
    curr_title = chapter_match.group(0) if chapter_match else "General"
    
    for i in range(1, len(segments)):
        seg = segments[i].strip()
        if not seg: continue
        lines = seg.split('\n')
        v_raw = " ".join([l.strip().lower() for l in lines[:3]])
        v_res = resolve(v_raw)
        
        record = {
            "book": current_book,
            "chapter_num": int(curr_ch),
            "chapter_title_ga": curr_title,
            "verse_num": int(v_res),
            "verse_ref": f"{current_book} {curr_ch}:{v_res}",
            "ga": seg,
            "en": "", # PDF usually missing interleaved English in this legacy layout
            "source_name": Path(file_path).name
        }
        node_text = BibleRefiner._build_bible_node_text(record)
        doc = Document(text=node_text, metadata={**record, "category": "bible", "file_path": file_path, "filename": Path(file_path).name})
        documents.append(doc)
    return documents


def _parse_bible_jsonl(text: str, file_path: str) -> List[Document]:
    """Parser for the structured Bible JSONL format."""
    documents = []
    for line in text.strip().split('\n'):
        if not line.strip(): continue
        try: record = json.loads(line)
        except: continue
        
        # Ensure chapter and verse are integers for strict filtering
        if "chapter_num" in record:
            try: record["chapter_num"] = int(record["chapter_num"])
            except: pass
        if "verse_num" in record:
            try: record["verse_num"] = int(record["verse_num"])
            except: pass
            
        # Force source_name to be the filename, not internal version metadata
        record["source_name"] = Path(file_path).name
        
        node_text = BibleRefiner._build_bible_node_text(record)
        doc = Document(text=node_text, metadata={**record, "category": "bible", "file_path": file_path, "filename": Path(file_path).name})
        documents.append(doc)
    return documents


def initial_index_if_needed() -> bool:
    """
    Perform initial indexing if collection does not exist.
    Only runs once if collection is empty.
    
    Returns:
        bool: True if collection was created, False if collection already exists or Qdrant unavailable
    """
    try:
        if collection_exists(COLLECTION_NAME):
            # Collection exists, skip indexing
            return False
        
        # Collection doesn't exist, but we don't have documents to index yet
        # This will be handled by the upload endpoint or worker
        # Just create the collection structure
        get_vector_store(COLLECTION_NAME)
        return True
    except (ConnectionError, Exception) as e:
        # Qdrant not available - this is OK, it will be available when needed
        # Log the issue but don't fail startup
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"Qdrant not available during startup: {e}. The application will start, but vector operations will fail until Qdrant is running.")
        return False

